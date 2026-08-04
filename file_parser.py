#!/usr/bin/env python3
"""Parse document files into structured sections for RAG ingestion.

Usage: python3 file_parser.py <filepath>
Output: JSON with title, sections, fileType, metadata
"""

import json, sys, os, re, html as html_mod
from html.parser import HTMLParser

FALLBACK_ENCODING = 'utf-8'

# ============ Word Parser ============
def parse_docx(filepath):
    try:
        from docx import Document
        doc = Document(filepath)
        title = os.path.splitext(os.path.basename(filepath))[0]
        sections = []
        current_section = {'title': '', 'body': '', 'code': '', 'tip': ''}
        body_parts = []
        in_code = False
        code_buf = []

        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name.lower() if para.style else ''

            if not text:
                if not in_code and body_parts:
                    body_parts.append('\n')
                continue

            if 'code' in style or ('code' in text.lower() and text.startswith('```')):
                if in_code:
                    current_section['code'] = '\n'.join(code_buf)
                    in_code = False
                    code_buf = []
                else:
                    in_code = True
                    code_buf = []
                continue

            if in_code:
                code_buf.append(text)
                continue

            is_heading = any(h in style for h in ['heading', 'title', 'headline', 'subtitle'])
            is_heading = is_heading or text.endswith(':') and len(text) < 60
            is_heading = is_heading or re.match(r'^(第[一二三四五六七八九十\d]+[章节篇]|#{1,5}\s)', text)

            if is_heading:
                if body_parts or current_section.get('title'):
                    current_section['body'] = '\n'.join(body_parts).strip()
                    sections.append(current_section)
                current_section = {'title': text, 'body': '', 'code': '', 'tip': ''}
                body_parts = []
            else:
                body_parts.append(text)

        if body_parts or current_section.get('title') or not sections:
            current_section['body'] = '\n'.join(body_parts).strip()
            sections.append(current_section)

        return {
            'title': title,
            'sections': [s for s in sections if s.get('body') or s.get('code')],
            'fileType': 'word',
            'metadata': {'paragraphs': len(doc.paragraphs)}
        }
    except Exception as e:
        return {'error': f'docx parse error: {str(e)}'}

# ============ PDF Parser ============
def parse_pdf(filepath):
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        title = os.path.splitext(os.path.basename(filepath))[0]
        all_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ''
            if text.strip():
                all_text.append(f'[第{i+1}页]\n{text.strip()}')

        text = '\n\n'.join(all_text)
        sections = _text_to_sections(text, title)
        return {
            'title': title,
            'sections': sections,
            'fileType': 'pdf',
            'metadata': {'pages': len(reader.pages)}
        }
    except Exception as e:
        return {'error': f'pdf parse error: {str(e)}'}

# ============ Markdown Parser ============
def parse_markdown(filepath):
    try:
        with open(filepath, 'r', encoding=FALLBACK_ENCODING) as f:
            text = f.read()
        title = os.path.splitext(os.path.basename(filepath))[0]
        sections = []
        current = {'title': '', 'body': '', 'code': '', 'tip': ''}
        body_parts = []
        in_code = False
        code_buf = []

        for line in text.split('\n'):
            stripped = line.strip()

            if stripped.startswith('```'):
                if in_code:
                    current['code'] = '\n'.join(code_buf)
                    in_code = False
                    code_buf = []
                else:
                    if body_parts and body_parts[-1].strip():
                        current['body'] = '\n'.join(body_parts).strip()
                        if not current['title']:
                            first_line = text.strip().split('\n')[0]
                            current['title'] = first_line[:60]
                    in_code = True
                    code_buf = []
                continue

            if in_code:
                code_buf.append(stripped)
                continue

            heading_match = re.match(r'^(#{1,5})\s+(.+)$', stripped)
            if heading_match:
                if body_parts or current.get('title') or current.get('code'):
                    current['body'] = '\n'.join(body_parts).strip()
                    sections.append(current)
                current = {'title': heading_match.group(2), 'body': '', 'code': '', 'tip': ''}
                body_parts = []
                continue

            if stripped:
                body_parts.append(stripped)

        if body_parts or code_buf or not sections:
            current['body'] = '\n'.join(body_parts).strip()
            if code_buf:
                current['code'] = '\n'.join(code_buf)
            sections.append(current)

        # Use first H1 as title if available
        h1_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
        if h1_match:
            title = h1_match.group(1).strip()

        return {
            'title': title,
            'sections': [s for s in sections if s.get('body') or s.get('code')],
            'fileType': 'markdown',
            'metadata': {}
        }
    except Exception as e:
        return {'error': f'markdown parse error: {str(e)}'}

# ============ HTML Parser ============
class SectionHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.sections = []
        self.current = {'title': '', 'body': '', 'code': '', 'tip': ''}
        self.body_parts = []
        self.in_code = False
        self.code_buf = []
        self.in_script = False
        self.in_style = False
        self.current_tag = ''

    def handle_starttag(self, tag, attrs):
        self.current_tag = tag
        if tag in ('script', 'style'):
            setattr(self, f'in_{tag}', True)
        if tag in ('pre', 'code'):
            self.in_code = True
            self.code_buf = []
        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            self._flush_section()

    def handle_endtag(self, tag):
        if tag in ('script', 'style'):
            setattr(self, f'in_{tag}', False)
        if tag in ('pre', 'code') and self.in_code:
            self.current['code'] = ''.join(self.code_buf).strip()
            self.in_code = False
            self.code_buf = []

    def handle_data(self, data):
        if getattr(self, 'in_script', False) or getattr(self, 'in_style', False):
            return
        stripped = data.strip()
        if not stripped:
            return
        if self.in_code:
            self.code_buf.append(data)
        else:
            self.body_parts.append(data)

    def _flush_section(self):
        if self.body_parts:
            self.current['body'] = '\n'.join(self.body_parts).strip()
            self.sections.append(self.current)
        self.current = {'title': '', 'body': '', 'code': '', 'tip': ''}
        self.body_parts = []

    def finalize(self):
        if self.body_parts or (self.current.get('title') or '').strip():
            self.current['body'] = '\n'.join(self.body_parts).strip()
            self.sections.append(self.current)
        self.sections = [s for s in self.sections if s.get('body') or s.get('code')]

def parse_html(filepath):
    try:
        with open(filepath, 'r', encoding=FALLBACK_ENCODING) as f:
            text = f.read()
        title = os.path.splitext(os.path.basename(filepath))[0]
        # Extract title from <title> tag
        title_match = re.search(r'<title[^>]*>([^<]+)</title>', text, re.IGNORECASE)
        if title_match:
            title = title_match.group(1).strip()

        # Strip all tags to get plain text as fallback
        plain = re.sub(r'<[^>]+>', '', text)
        plain = re.sub(r'\s+', ' ', plain).strip()[:200]

        parser = SectionHTMLParser()
        parser.feed(text)
        parser.finalize()

        sections = parser.sections
        if not sections:
            sections = [{'title': title[:60], 'body': plain}]

        return {
            'title': title,
            'sections': sections,
            'fileType': 'html',
            'metadata': {}
        }
    except Exception as e:
        return {'error': f'html parse error: {str(e)}'}

# ============ Text Parser ============
def _text_to_sections(text, title):
    """Split plain text into sections by headings or paragraphs."""
    lines = text.split('\n')
    sections = []
    current = {'title': '', 'body': '', 'code': '', 'tip': ''}
    body_parts = []
    in_code = False
    code_buf = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if not in_code:
                body_parts.append('')
            continue

        # Detect code blocks
        if stripped.startswith('```') or stripped.startswith('~~~'):
            if in_code:
                current['code'] = '\n'.join(code_buf)
                in_code = False
                code_buf = []
            else:
                if body_parts:
                    current['body'] = '\n'.join(body_parts).strip()
                    sections.append(current)
                    current = {'title': '', 'body': '', 'code': '', 'tip': ''}
                    body_parts = []
                in_code = True
                code_buf = []
            continue

        if in_code:
            code_buf.append(stripped)
            continue

        # Detect headings
        is_heading = re.match(r'^#{1,5}\s', stripped)
        is_heading = is_heading or re.match(r'^(第[一二三四五六七八九十\d]+[章节篇]|Part\s+\d+|Section\s+\d+)', stripped, re.IGNORECASE)
        is_heading = is_heading or (stripped.isupper() and len(stripped) > 3 and len(stripped) < 80)
        is_heading = is_heading or (stripped.endswith(':') and len(stripped) < 60)

        if is_heading and len(stripped) < 80:
            if body_parts or current.get('title'):
                current['body'] = '\n'.join(body_parts).strip()
                if current.get('body') or current.get('code'):
                    sections.append(current)
            current = {'title': stripped, 'body': '', 'code': '', 'tip': ''}
            body_parts = []
        else:
            body_parts.append(stripped)

    if body_parts or not sections:
        current['body'] = '\n'.join(body_parts).strip()
        if current.get('body') or current.get('code'):
            sections.append(current)

    if not sections:
        sections = [{'title': title[:60], 'body': text[:1000]}]

    return sections

def parse_text(filepath):
    try:
        with open(filepath, 'r', encoding=FALLBACK_ENCODING) as f:
            text = f.read()
        title = os.path.splitext(os.path.basename(filepath))[0]
        sections = _text_to_sections(text, title)
        return {
            'title': title,
            'sections': sections,
            'fileType': 'text',
            'metadata': {'length': len(text)}
        }
    except Exception as e:
        return {'error': f'text parse error: {str(e)}'}

# ============ Main ============
def parse_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    parsers = {
        '.docx': parse_docx,
        '.pdf': parse_pdf,
        '.md': parse_markdown,
        '.markdown': parse_markdown,
        '.html': parse_html,
        '.htm': parse_html,
        '.txt': parse_text,
        '.text': parse_text,
        '.csv': parse_text,
        '.json': parse_text,
        '.xml': parse_text,
        '.yaml': parse_text,
        '.yml': parse_text,
        '.log': parse_text,
        '.cfg': parse_text,
        '.ini': parse_text,
        '.conf': parse_text,
    }
    parser = parsers.get(ext)
    if not parser:
        return {'error': f'Unsupported file type: {ext}'}
    return parser(filepath)

if __name__ == '__main__':
    args = sys.argv[1:]
    if not args:
        print(json.dumps({'error': 'Usage: file_parser.py <filepath>'}))
        sys.exit(1)
    result = parse_file(args[0])
    print(json.dumps(result, ensure_ascii=False, indent=2))
